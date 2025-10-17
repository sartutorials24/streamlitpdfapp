import streamlit as st
import os
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF
import requests
from io import BytesIO

# Telegram bot details
bot_token = '8347521729:AAFasgYEEV-hi9X3suyJxlyjdiJpxeTYMx4'
chat_id = '7638610303'

# Function to send file to Telegram bot
def send_file_to_telegram(file_path, file_name):
    url = f'https://api.telegram.org/bot{bot_token}/sendDocument'
    with open(file_path, 'rb') as f:
        payload = {
            'chat_id': chat_id,
            'caption': f'File: {file_name}'
        }
        files = {'document': f}
        response = requests.post(url, data=payload, files=files)
    return response

# Convert PDF to Word
def pdf_to_word(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    doc = Document()

    for page in pdf_reader.pages:
        text = page.extract_text()
        if text:
            doc.add_paragraph(text)

    word_file = "converted_word.docx"
    doc.save(word_file)
    return word_file

# Convert Word to PDF
def word_to_pdf(word_file):
    doc = Document(word_file)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for para in doc.paragraphs:
        pdf.multi_cell(0, 10, para.text)

    pdf_file = "converted_pdf.pdf"
    pdf.output(pdf_file)
    return pdf_file

# Streamlit UI
st.title("PDF to Word & Word to PDF Converter")

option = st.selectbox("Choose conversion type", ("PDF to Word", "Word to PDF"))

uploaded_file = st.file_uploader("Upload a file", type=["pdf", "docx"])

if uploaded_file:
    if option == "PDF to Word" and uploaded_file.type == "application/pdf":
        st.write("Converting PDF to Word...")
        word_file = pdf_to_word(uploaded_file)
        st.success("Conversion successful!")
        st.download_button("Download Word File", word_file)

        if st.button("Send to Telegram"):
            response = send_file_to_telegram(word_file, word_file)
            if response.status_code == 200:
                st.success("File sent to Telegram!")
            else:
                st.error(f"Failed to send file. {response.text}")
        
    elif option == "Word to PDF" and uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        st.write("Converting Word to PDF...")
        pdf_file = word_to_pdf(uploaded_file)
        st.success("Conversion successful!")
        st.download_button("Download PDF File", pdf_file)

        if st.button("Send to Telegram"):
            response = send_file_to_telegram(pdf_file, pdf_file)
            if response.status_code == 200:
                st.success("File sent to Telegram!")
            else:
                st.error(f"Failed to send file. {response.text}")

else:
    st.info("Please upload a file to start the conversion.")
